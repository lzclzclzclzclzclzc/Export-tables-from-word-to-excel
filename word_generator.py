from docx import Document

def create_table(doc, rows, cols, table_num):
    """在文档中插入一个指定行列的表格，并合并部分单元格"""
    doc.add_heading(f'表格 {table_num}', level=2)
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'

    # 填入内容
    for r in range(rows):
        for c in range(cols):
            table.cell(r, c).text = f'R{r+1}C{c+1}'

    # # 安全合并一些单元格
    # if cols >= 2:
    #     table.cell(0, 0).merge(table.cell(0, 1))  # 合并第一行前两列
    # if rows >= 3:
    #     table.cell(1, 0).merge(table.cell(2, 0))  # 合并第一列中第二三行
    # if cols >= 4:
    #     table.cell(0, 2).merge(table.cell(0, 3))  # 合并第一行中间两列

def generate_doc(filename="input_docx.docx"):
    """创建包含10个表格的 Word 文档"""
    doc = Document()
    doc.add_heading('包含合并单元格的10个表格示例', level=1)

    # 每个表格的大小（行, 列）
    table_sizes = [(3, 4), (5, 3), (4, 5), (6, 2), (2, 6),
                   (4, 4), (5, 5), (3, 3), (6, 4), (4, 6)]

    for i, (rows, cols) in enumerate(table_sizes, start=1):
        create_table(doc, rows, cols, i)

    doc.save(filename)
    print(f"文档已保存为：{filename}")

if __name__ == "__main__":
    generate_doc()