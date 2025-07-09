from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment
import os
import glob
import sys
import re
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
from tkinter.font import Font
import threading


def extract_tables_with_formatting(docx_path):
    """提取Word表格并保留基础格式和多行文本"""
    doc = Document(docx_path)
    tables_data = []

    for table_idx, table in enumerate(doc.tables, 1):
        table_data = []
        max_cols = 0

        # 确定最大列数
        for row in table.rows:
            col_count = len(row.cells)
            if col_count > max_cols:
                max_cols = col_count

        # 提取表格数据
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                # 处理单元格文本，保留换行符
                cell_text = ""
                for paragraph in cell.paragraphs:
                    # 保留段落内的换行符
                    for run in paragraph.runs:
                        # 保留原始换行符
                        cell_text += run.text

                    # 段落之间添加换行符（除非是最后一个段落）
                    if paragraph != cell.paragraphs[-1]:
                        cell_text += "\n"

                # 清理多余的空格但保留换行
                cell_text = re.sub(r'[ \t]+', ' ', cell_text)  # 压缩连续空格
                cell_text = re.sub(r'\n\s+', '\n', cell_text)  # 清理行首空格
                cell_text = cell_text.strip()

                # 获取对齐方式
                alignment = None
                if cell.paragraphs:
                    para = cell.paragraphs[0]
                    if hasattr(para, 'alignment') and para.alignment:
                        alignment = str(para.alignment)

                # 添加到行数据
                row_data.append({
                    'text': cell_text,
                    'alignment': alignment,
                    'colspan': 1  # 默认列跨度
                })

            # 填充缺失的单元格（确保每行列数一致）
            while len(row_data) < max_cols:
                row_data.append({
                    'text': "",
                    'alignment': None,
                    'colspan': 1
                })

            table_data.append(row_data)

        # 添加到表格数据
        tables_data.append({
            'index': table_idx,
            'data': table_data
        })

    return tables_data


def export_to_excel(tables_data, output_path):
    """将表格数据导出到Excel的不同Sheet，正确处理多行文本"""
    wb = Workbook()
    # 删除默认创建的空Sheet
    if 'Sheet' in wb.sheetnames:
        wb.remove(wb['Sheet'])

    if not tables_data:
        # 没有表格时创建提示Sheet
        ws = wb.create_sheet(title="No Tables")
        ws['A1'] = "未在文档中找到任何表格"
        wb.save(output_path)
        return

    for table in tables_data:
        # 创建Sheet（名称长度限制31字符）
        sheet_name = f"Table_{table['index']}"[:31]
        ws = wb.create_sheet(title=sheet_name)

        table_data = table['data']

        # 写入数据并应用格式
        for row_idx, row_data in enumerate(table_data, 1):
            for col_idx, cell_info in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=cell_info['text'])

                # 应用文本对齐方式
                align_obj = Alignment(vertical='center')  # 默认垂直居中

                # 设置水平对齐
                if cell_info['alignment']:
                    if 'CENTER' in cell_info['alignment']:
                        align_obj.horizontal = 'center'
                    elif 'RIGHT' in cell_info['alignment']:
                        align_obj.horizontal = 'right'
                    elif 'LEFT' in cell_info['alignment']:
                        align_obj.horizontal = 'left'

                # 如果文本中有换行符，设置自动换行
                if cell_info['text'] and '\n' in cell_info['text']:
                    align_obj.wrapText = True

                # 应用对齐样式
                cell.alignment = align_obj

    # 自动调整列宽
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        for col in ws.columns:
            max_length = 0
            column = col[0].column_letter  # 获取列字母

            for cell in col:
                try:
                    # 计算多行文本的最大长度
                    if cell.value:
                        # 对于多行文本，取最长行的长度
                        lines = str(cell.value).split('\n')
                        max_line_length = max(len(line) for line in lines)
                        if max_line_length > max_length:
                            max_length = max_line_length
                except:
                    pass

            # 设置列宽（加一点缓冲空间）
            if max_length > 0:
                adjusted_width = (max_length + 2) * 1.2
                ws.column_dimensions[column].width = min(adjusted_width, 50)
            else:
                ws.column_dimensions[column].width = 10  # 默认宽度

    wb.save(output_path)


class WordTableExtractorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Word表格提取工具")
        self.root.geometry("800x600")
        self.root.resizable(True, True)

        # 设置应用图标
        try:
            self.root.iconbitmap("icon.ico")
        except:
            pass

        # 创建主框架
        self.main_frame = ttk.Frame(root, padding="20")
        self.main_frame.pack(fill=tk.BOTH, expand=True)

        # 标题
        title_font = Font(family="微软雅黑", size=16, weight="bold")
        self.title_label = ttk.Label(
            self.main_frame,
            text="Word文档表格提取工具",
            font=title_font,
            anchor=tk.CENTER
        )
        self.title_label.pack(pady=(0, 20))

        # 文件夹选择区域
        folder_frame = ttk.LabelFrame(self.main_frame, text="选择文件夹")
        folder_frame.pack(fill=tk.X, padx=5, pady=5)

        self.folder_path = tk.StringVar(value=os.getcwd())

        folder_entry = ttk.Entry(folder_frame, textvariable=self.folder_path, width=60)
        folder_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 5))

        browse_btn = ttk.Button(
            folder_frame,
            text="浏览...",
            width=10,
            command=self.browse_folder
        )
        browse_btn.pack(side=tk.RIGHT)

        # 按钮区域
        btn_frame = ttk.Frame(self.main_frame)
        btn_frame.pack(fill=tk.X, pady=10)

        self.extract_btn = ttk.Button(
            btn_frame,
            text="提取表格",
            command=self.start_extraction,
            width=15
        )
        self.extract_btn.pack(side=tk.LEFT, padx=5)

        self.open_folder_btn = ttk.Button(
            btn_frame,
            text="打开输出文件夹",
            command=self.open_output_folder,
            width=15,
            state=tk.DISABLED
        )
        self.open_folder_btn.pack(side=tk.LEFT, padx=5)

        self.clear_btn = ttk.Button(
            btn_frame,
            text="清空日志",
            command=self.clear_log,
            width=15
        )
        self.clear_btn.pack(side=tk.RIGHT, padx=5)

        # 日志区域
        log_frame = ttk.LabelFrame(self.main_frame, text="处理日志")
        log_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        self.log_text = scrolledtext.ScrolledText(
            log_frame,
            wrap=tk.WORD,
            font=("Consolas", 10)
        )
        self.log_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        self.log_text.config(state=tk.DISABLED)

        # 状态栏
        self.status_var = tk.StringVar(value="就绪")
        status_bar = ttk.Label(
            root,
            textvariable=self.status_var,
            relief=tk.SUNKEN,
            anchor=tk.W
        )
        status_bar.pack(side=tk.BOTTOM, fill=tk.X)

        # 处理统计
        self.success_count = 0
        self.error_count = 0
        self.total_files = 0

        # 设置样式
        self.style = ttk.Style()
        self.style.configure("TButton", padding=6)
        self.style.configure("TLabelFrame", padding=10)

    def browse_folder(self):
        """打开文件夹选择对话框"""
        folder_selected = filedialog.askdirectory(
            initialdir=self.folder_path.get(),
            title="选择包含Word文档的文件夹"
        )
        if folder_selected:
            self.folder_path.set(folder_selected)

    def log_message(self, message):
        """向日志区域添加消息"""
        self.log_text.config(state=tk.NORMAL)
        self.log_text.insert(tk.END, message + "\n")
        self.log_text.config(state=tk.DISABLED)
        self.log_text.see(tk.END)  # 滚动到底部

    def clear_log(self):
        """清空日志区域"""
        self.log_text.config(state=tk.NORMAL)
        self.log_text.delete(1.0, tk.END)
        self.log_text.config(state=tk.DISABLED)
        self.status_var.set("日志已清空")

    def open_output_folder(self):
        """打开输出文件夹"""
        output_folder = self.folder_path.get()
        if os.path.exists(output_folder):
            if sys.platform == "win32":
                os.startfile(output_folder)
            elif sys.platform == "darwin":
                os.system(f"open '{output_folder}'")
            else:
                os.system(f"xdg-open '{output_folder}'")

    def start_extraction(self):
        """开始提取过程"""
        directory = self.folder_path.get()

        if not os.path.exists(directory):
            messagebox.showerror("错误", "指定的文件夹不存在！")
            return

        # 禁用按钮
        self.extract_btn.config(state=tk.DISABLED)
        self.open_folder_btn.config(state=tk.DISABLED)

        # 重置统计
        self.success_count = 0
        self.error_count = 0

        # 清空日志
        self.clear_log()
        self.log_message(f"开始处理文件夹: {directory}")
        self.log_message("-" * 70)

        # 在后台线程中运行处理
        threading.Thread(
            target=self.process_docx_files,
            args=(directory,),
            daemon=True
        ).start()

    def process_docx_files(self, directory):
        """处理指定目录中的所有Word文档"""
        # 查找目录中的所有docx文件
        docx_files = glob.glob(os.path.join(directory, "*.docx"))
        self.total_files = len(docx_files)

        if not docx_files:
            self.log_message("当前目录中没有找到任何Word文档(.docx文件)")
            self.status_var.set("处理完成: 未找到Word文档")
            self.extract_btn.config(state=tk.NORMAL)
            return

        self.log_message(f"找到 {self.total_files} 个Word文档，开始处理...")
        self.log_message("=" * 70)

        for i, docx_path in enumerate(docx_files, 1):
            try:
                filename = os.path.basename(docx_path)
                self.log_message(f"正在处理文件 ({i}/{self.total_files}): {filename}")

                # 生成输出文件名（相同文件名，扩展名改为.xlsx）
                excel_filename = os.path.splitext(filename)[0] + "_表格导出.xlsx"
                output_path = os.path.join(directory, excel_filename)

                # 提取表格并导出
                tables = extract_tables_with_formatting(docx_path)
                export_to_excel(tables, output_path)

                self.log_message(f"  成功导出: {excel_filename}")
                self.log_message(f"  找到表格数量: {len(tables)}")
                self.success_count += 1
                self.log_message("-" * 70)

            except Exception as e:
                self.log_message(f"  处理出错: {str(e)}")
                self.log_message("-" * 70)
                self.error_count += 1
                continue

        # 处理结果汇总
        self.log_message("=" * 70)
        result_msg = f"处理完成! 成功: {self.success_count}/{self.total_files}, 失败: {self.error_count}/{self.total_files}"
        self.log_message(result_msg)

        if self.error_count > 0:
            self.log_message("注意：部分文件处理失败，请查看错误信息")

        # 更新状态
        self.status_var.set(result_msg)

        # 启用按钮
        self.extract_btn.config(state=tk.NORMAL)
        self.open_folder_btn.config(state=tk.NORMAL)


if __name__ == "__main__":
    root = tk.Tk()
    app = WordTableExtractorApp(root)
    root.mainloop()